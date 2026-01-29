"""Curriculum import parsing and insertion logic."""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
import shutil
import subprocess
from typing import Dict, List, Optional, Tuple

from ..models.database import Database
from ..models.entities import Teacher


@dataclass
class CurriculumQuestion:
    """Parsed question item."""
    number: Optional[int]
    text: str


@dataclass
class CurriculumLesson:
    """Parsed lesson with hours and questions."""
    number: Optional[int]
    title: str
    lesson_type_name: Optional[str]
    total_hours: Optional[float]
    classroom_hours: Optional[float]
    self_study_hours: Optional[float]
    questions: List[CurriculumQuestion] = field(default_factory=list)


@dataclass
class CurriculumTopic:
    """Parsed topic with lessons."""
    title: str
    lessons: List[CurriculumLesson] = field(default_factory=list)


LESSON_TYPE_KEYWORDS = [
    ("контрольне заняття", "Контрольне заняття"),
    ("групове заняття", "Групове заняття"),
    ("практичне заняття", "Практичне заняття"),
    ("семінар", "Семінар"),
    ("лекція", "Лекція"),
]


def parse_curriculum_text(text: str) -> List[CurriculumTopic]:
    """
    Parse a structured curriculum table pasted from Word/Excel/text.

    Mapping rules:
    - Topic title from 'Назва теми'
    - Lesson title from 'Заняття X. Назва'
    - Lesson type detected by keywords in lesson title
    - Hours from total/classroom/self-study columns (if present)
    - Each numbered line under lesson is a separate question
    """
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    if not lines:
        raise ValueError("Empty input.")

    delimiter = _detect_delimiter(lines)
    rows = [_split_row(line, delimiter) for line in lines]

    header_idx = _find_header_row(rows)
    if header_idx is None:
        # Fallback: parse free-form text with "Тема X. Назва" and "Заняття X. Назва"
        return _parse_freeform_curriculum(text)

    headers = [_normalize_header(cell) for cell in rows[header_idx]]
    col_map = _map_columns(headers)
    if col_map["topic"] is None or col_map["lesson"] is None:
        raise ValueError("Required columns missing: 'Назва теми' and 'Заняття'.")

    topics: List[CurriculumTopic] = []
    current_topic: Optional[CurriculumTopic] = None
    lesson_auto_index = 1

    for row in rows[header_idx + 1:]:
        if _is_summary_row(row):
            continue

        topic_cell = _cell(row, col_map["topic"])
        lesson_cell = _cell(row, col_map["lesson"])
        if topic_cell:
            current_topic = CurriculumTopic(title=_normalize_text(topic_cell))
            topics.append(current_topic)
            lesson_auto_index = 1

        if not current_topic:
            continue

        if not lesson_cell:
            continue

        lesson_number, lesson_title, lesson_questions = _parse_lesson_cell(lesson_cell)
        if lesson_number is None:
            lesson_number = lesson_auto_index
        lesson_auto_index = lesson_number + 1

        lesson_type_name = _detect_lesson_type(lesson_title)

        total_hours = _parse_number(_cell(row, col_map["total_hours"]))
        classroom_hours = _parse_number(_cell(row, col_map["classroom_hours"]))
        self_study_hours = _parse_number(_cell(row, col_map["self_study_hours"]))

        if total_hours is None and classroom_hours is not None and self_study_hours is not None:
            total_hours = classroom_hours + self_study_hours

        questions_text = _cell(row, col_map["questions"])
        if questions_text:
            lesson_questions.extend(_parse_questions_block(questions_text))

        lesson = CurriculumLesson(
            number=lesson_number,
            title=_normalize_text(lesson_title),
            lesson_type_name=lesson_type_name,
            total_hours=total_hours,
            classroom_hours=classroom_hours,
            self_study_hours=self_study_hours,
            questions=lesson_questions,
        )
        current_topic.lessons.append(lesson)

    return topics


def _parse_freeform_curriculum(text: str) -> List[CurriculumTopic]:
    topics: List[CurriculumTopic] = []
    current_topic: Optional[CurriculumTopic] = None
    current_lesson: Optional[CurriculumLesson] = None

    prepared = _prepare_freeform_text(text)
    for raw_line in prepared.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Ignore summary lines
        if _is_summary_row([line]):
            continue

        topic_match = re.match(r"тема\s*(\d+)\.?\s*(.+)", line, flags=re.IGNORECASE)
        if topic_match:
            title = _normalize_text(line)
            current_topic = CurriculumTopic(title=title)
            topics.append(current_topic)
            current_lesson = None
            continue

        lesson_number, lesson_title, lesson_questions = _parse_lesson_cell(line)
        if lesson_title and re.search(r"заняття\s*\d+", line, flags=re.IGNORECASE):
            lesson_type_name = _detect_lesson_type(lesson_title)
            inline_questions = _parse_questions_block(line)
            if inline_questions:
                lesson_questions.extend(inline_questions)
            current_lesson = CurriculumLesson(
                number=lesson_number,
                title=_normalize_text(lesson_title),
                lesson_type_name=lesson_type_name,
                total_hours=None,
                classroom_hours=None,
                self_study_hours=None,
                questions=lesson_questions,
            )
            if current_topic is None:
                current_topic = CurriculumTopic(title="Тема")
                topics.append(current_topic)
            current_topic.lessons.append(current_lesson)
            continue

        # Questions are lines after "Заняття X. ..." until next lesson/topic
        if current_lesson:
            current_lesson.questions.extend(_parse_questions_block(line))

    return topics


def extract_text_from_file(path: str) -> str:
    """Extract text from .txt/.csv/.tsv/.docx/.doc files for parsing."""
    ext = Path(path).suffix.lower()
    if ext in {".txt", ".csv", ".tsv"}:
        with open(path, "r", encoding="utf-8") as file:
            return file.read()
    if ext == ".docx":
        return _extract_from_docx(path)
    if ext == ".doc":
        return _extract_from_doc(path)
    raise ValueError("Unsupported file type.")


def parse_teachers_from_docx(path: str) -> List[Teacher]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx is required to import .docx files.") from exc

    doc = Document(path)
    if not doc.tables:
        raise ValueError("No tables found in .docx file.")

    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        raise ValueError("Empty table in .docx file.")

    header = [cell.lower() for cell in rows[0]]
    has_header = any("п" in cell or "name" in cell or "email" in cell for cell in header)
    start_idx = 1 if has_header else 0
    col_map = _map_teacher_columns(header) if has_header else {}

    teachers: List[Teacher] = []
    for row in rows[start_idx:]:
        def get(idx_name, fallback_idx):
            if idx_name in col_map:
                idx = col_map[idx_name]
                return row[idx].strip() if idx < len(row) else ""
            if fallback_idx is not None and fallback_idx < len(row):
                return row[fallback_idx].strip()
            return ""

        full_name = get("full_name", 0)
        if not full_name:
            continue
        teacher = Teacher(
            full_name=full_name,
            military_rank=get("military_rank", 1) or None,
            position=get("position", 2) or None,
            department=get("department", 3) or None,
            email=get("email", 4) or None,
            phone=get("phone", 5) or None,
        )
        teachers.append(teacher)

    return teachers


def import_teachers_from_docx(database: Database, path: str) -> Tuple[int, int]:
    """Import teachers from .docx table. Returns (added, skipped)."""
    teachers = parse_teachers_from_docx(path)
    if not teachers:
        raise ValueError("No teachers found in .docx table.")

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT full_name, email FROM teachers")
        existing = set()
        for row in cursor.fetchall():
            name = (row["full_name"] or "").lower()
            email = (row["email"] or "").lower()
            if name:
                existing.add(name)
            if email:
                existing.add(email)

        added = 0
        skipped = 0
        for teacher in teachers:
            key_name = (teacher.full_name or "").lower()
            key_email = (teacher.email or "").lower()
            if key_name in existing or (key_email and key_email in existing):
                skipped += 1
                continue
            cursor.execute("""
                INSERT INTO teachers (full_name, military_rank, position, department, email, phone)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                teacher.full_name,
                teacher.military_rank,
                teacher.position,
                teacher.department,
                teacher.email,
                teacher.phone,
            ))
            added += 1
            existing.add(key_name)
            if key_email:
                existing.add(key_email)

    return added, skipped


def import_curriculum_structure(
    database: Database,
    program_id: int,
    discipline_id: Optional[int],
    new_discipline_name: Optional[str],
    topics: List[CurriculumTopic],
) -> Tuple[int, int, int]:
    """
    Import parsed curriculum into the database.

    Returns:
        Tuple of (topics_added, lessons_added, questions_added).
    """
    if not topics:
        raise ValueError("No topics to import.")

    with database.get_connection() as conn:
        cursor = conn.cursor()
        discipline_id = _ensure_discipline(cursor, program_id, discipline_id, new_discipline_name)
        return _import_with_cursor(cursor, discipline_id, topics)


def import_curriculum_structure_by_names(
    database: Database,
    program_name: str,
    discipline_name: str,
    topics: List[CurriculumTopic],
) -> Tuple[int, int, int]:
    if not program_name:
        raise ValueError("Program name is required.")
    if not discipline_name:
        raise ValueError("Discipline name is required.")
    with database.get_connection() as conn:
        cursor = conn.cursor()
        program_id = _ensure_program(cursor, program_name)
        discipline_id = _ensure_discipline_by_name(cursor, program_id, discipline_name)
        return _import_with_cursor(cursor, discipline_id, topics)


def program_discipline_from_filename(path: str) -> Tuple[str, Optional[str]]:
    name = Path(path).stem
    if " - " in name:
        parts = [part.strip() for part in name.split(" - ", 1)]
        return parts[0], parts[1] if len(parts) > 1 else None
    return name, None


def _import_with_cursor(cursor, discipline_id: int, topics: List[CurriculumTopic]) -> Tuple[int, int, int]:
    lesson_type_map = _ensure_lesson_types(cursor)

    existing_topics = _load_existing_topics(cursor, discipline_id)
    existing_lessons = _load_existing_lessons(cursor, discipline_id)

    topics_added = 0
    lessons_added = 0
    questions_added = 0

    for topic_index, topic in enumerate(topics, start=1):
        topic_key = _key(topic.title)
        if topic_key in existing_topics:
            topic_id = existing_topics[topic_key]
        else:
            cursor.execute(
                "INSERT INTO topics (title, description, order_index) VALUES (?, ?, ?)",
                (topic.title, "", topic_index),
            )
            topic_id = cursor.lastrowid
            cursor.execute(
                "INSERT OR IGNORE INTO discipline_topics (discipline_id, topic_id, order_index) VALUES (?, ?, ?)",
                (discipline_id, topic_id, topic_index),
            )
            existing_topics[topic_key] = topic_id
            topics_added += 1

        lesson_order = _get_next_order_index(cursor, "topic_lessons", "topic_id", topic_id)
        for lesson in topic.lessons:
            lesson_key = (topic_id, _key(lesson.title))
            if lesson_key in existing_lessons:
                lesson_id, lesson_data = existing_lessons[lesson_key]
                _update_lesson_if_missing(cursor, lesson_id, lesson, lesson_type_map, lesson_data)
            else:
                lesson_type_id = lesson_type_map.get(lesson.lesson_type_name)
                cursor.execute(
                    """
                    INSERT INTO lessons (
                        title, description, duration_hours, lesson_type_id,
                        classroom_hours, self_study_hours, order_index
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        lesson.title,
                        "",
                        lesson.total_hours or 0.0,
                        lesson_type_id,
                        lesson.classroom_hours,
                        lesson.self_study_hours,
                        lesson.number or lesson_order,
                    ),
                )
                lesson_id = cursor.lastrowid
                cursor.execute(
                    "INSERT OR IGNORE INTO topic_lessons (topic_id, lesson_id, order_index) VALUES (?, ?, ?)",
                    (topic_id, lesson_id, lesson.number or lesson_order),
                )
                existing_lessons[lesson_key] = (
                    lesson_id,
                    {
                        "lesson_type_id": lesson_type_id,
                        "duration_hours": lesson.total_hours or 0.0,
                        "classroom_hours": lesson.classroom_hours,
                        "self_study_hours": lesson.self_study_hours,
                    },
                )
                lessons_added += 1
            lesson_order += 1

            questions_added += _insert_questions(cursor, lesson_id, lesson.questions)

    return topics_added, lessons_added, questions_added


def _detect_delimiter(lines: List[str]) -> str:
    if any("\t" in line for line in lines):
        return "\t"
    if any("|" in line for line in lines):
        return "|"
    return "  "


def _extract_from_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx is required to import .docx files.") from exc

    doc = Document(path)
    lines: List[str] = []
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append("\t".join(cells))
    else:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
    return "\n".join(lines)


def _extract_from_doc(path: str) -> str:
    tool = shutil.which("antiword") or shutil.which("catdoc")
    if not tool:
        raise ValueError("Importing .doc requires 'antiword' or 'catdoc' to be installed.")
    result = subprocess.run([tool, path], capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise ValueError("Failed to extract .doc text.")
    return result.stdout


def _split_row(line: str, delimiter: str) -> List[str]:
    if delimiter == "\t":
        return [cell.strip() for cell in line.split("\t")]
    if delimiter == "|":
        return [cell.strip() for cell in line.split("|")]
    return [cell.strip() for cell in re.split(r"\s{2,}", line.strip())]


def _find_header_row(rows: List[List[str]]) -> Optional[int]:
    for idx, row in enumerate(rows):
        joined = " ".join(row).lower()
        if "назва теми" in joined and "заняття" in joined:
            return idx
    return None


def _normalize_header(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _map_columns(headers: List[str]) -> Dict[str, Optional[int]]:
    col_map = {
        "topic": None,
        "lesson": None,
        "total_hours": None,
        "classroom_hours": None,
        "self_study_hours": None,
        "questions": None,
    }
    for idx, header in enumerate(headers):
        if "назва теми" in header:
            col_map["topic"] = idx
        elif "заняття" in header:
            col_map["lesson"] = idx
        elif "всього" in header or "усього" in header or "заг" in header:
            col_map["total_hours"] = idx
        elif "аудитор" in header or "лекц" in header or "практ" in header:
            col_map["classroom_hours"] = idx
        elif "самост" in header:
            col_map["self_study_hours"] = idx
        elif re.search(r"\bпитан(ня|ь)\b", header):
            col_map["questions"] = idx
    return col_map


def _map_teacher_columns(headers: List[str]) -> Dict[str, int]:
    mapping = {}
    for idx, header in enumerate(headers):
        h = _normalize_header(header)
        if "піб" in h or "full name" in h or "name" == h or "пріз" in h:
            mapping["full_name"] = idx
        elif "військ" in h or "зван" in h or "rank" in h:
            mapping["military_rank"] = idx
        elif "посад" in h or "position" in h:
            mapping["position"] = idx
        elif "каф" in h or "відділ" in h or "department" in h:
            mapping["department"] = idx
        elif "email" in h or "e-mail" in h or "пошта" in h:
            mapping["email"] = idx
        elif "тел" in h or "phone" in h:
            mapping["phone"] = idx
    return mapping


def _cell(row: List[str], idx: Optional[int]) -> str:
    if idx is None or idx >= len(row):
        return ""
    return row[idx].strip()


def _is_summary_row(row: List[str]) -> bool:
    joined = " ".join(row).lower()
    return "всього за тему" in joined or joined.startswith("всього")


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _key(text: str) -> str:
    return _normalize_text(text).lower()


def _detect_lesson_type(text: str) -> Optional[str]:
    lowered = text.lower()
    for keyword, name in LESSON_TYPE_KEYWORDS:
        if keyword in lowered:
            return name
    return None


def _parse_number(text: str) -> Optional[float]:
    if not text:
        return None
    match = re.search(r"[-+]?\d+(?:[.,]\d+)?", text.replace(" ", ""))
    if not match:
        return None
    value = match.group(0).replace(",", ".")
    try:
        return float(value)
    except ValueError:
        return None


def _parse_lesson_cell(text: str) -> Tuple[Optional[int], str, List[CurriculumQuestion]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    first_line = lines[0] if lines else text.strip()
    number = None
    title = first_line

    match = re.search(r"заняття\s*(\d+)\.?\s*(.*)", first_line, flags=re.IGNORECASE)
    if match:
        number = int(match.group(1))
        title = first_line
    else:
        num_match = re.match(r"(\d+)[).]\s*(.*)", first_line)
        if num_match:
            number = int(num_match.group(1))
            title = first_line

    questions = []
    if len(lines) > 1:
        questions.extend(_parse_questions_block("\n".join(lines[1:])))
    return number, title, questions


def _parse_questions_block(text: str) -> List[CurriculumQuestion]:
    questions: List[CurriculumQuestion] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        matches = list(re.finditer(r"(\d+)[).]\s*([^\\n]+?)(?=(\d+[).]|$))", line))
        if matches:
            for match in matches:
                num = int(match.group(1))
                questions.append(
                    CurriculumQuestion(number=num, text=_normalize_text(f"{num}. {match.group(2)}"))
                )
            continue
        match = re.match(r"(\d+)[).]\s*(.+)", line)
        if match:
            num = int(match.group(1))
            questions.append(CurriculumQuestion(number=num, text=_normalize_text(f"{num}. {match.group(2)}")))
    return questions


def _prepare_freeform_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"(?i)(?<!\n)(тема\s*\d+\.?)", r"\n\1", normalized)
    normalized = re.sub(r"(?i)(?<!\n)(заняття\s*\d+\.?)", r"\n\1", normalized)
    return normalized


def _ensure_discipline(cursor, program_id: int, discipline_id: Optional[int], new_name: Optional[str]) -> int:
    if discipline_id is None:
        if not new_name:
            raise ValueError("Discipline is required.")
        cursor.execute(
            "INSERT INTO disciplines (name, description, order_index) VALUES (?, ?, ?)",
            (new_name.strip(), "", _get_next_order_index(cursor, "program_disciplines", "program_id", program_id)),
        )
        discipline_id = cursor.lastrowid
    cursor.execute(
        "INSERT OR IGNORE INTO program_disciplines (program_id, discipline_id, order_index) VALUES (?, ?, ?)",
        (program_id, discipline_id, _get_next_order_index(cursor, "program_disciplines", "program_id", program_id)),
    )
    return discipline_id


def _ensure_program(cursor, program_name: str) -> int:
    cursor.execute("SELECT id FROM educational_programs WHERE name = ?", (program_name,))
    row = cursor.fetchone()
    if row:
        return row["id"]
    cursor.execute(
        """
        INSERT INTO educational_programs (name, description, level, duration_hours)
        VALUES (?, ?, ?, ?)
        """,
        (program_name, "", "", 0),
    )
    return cursor.lastrowid


def _ensure_discipline_by_name(cursor, program_id: int, discipline_name: str) -> int:
    cursor.execute(
        """
        SELECT d.id
        FROM disciplines d
        JOIN program_disciplines pd ON d.id = pd.discipline_id
        WHERE pd.program_id = ? AND d.name = ?
        """,
        (program_id, discipline_name),
    )
    row = cursor.fetchone()
    if row:
        discipline_id = row["id"]
    else:
        cursor.execute(
            "INSERT INTO disciplines (name, description, order_index) VALUES (?, ?, ?)",
            (discipline_name, "", _get_next_order_index(cursor, "program_disciplines", "program_id", program_id)),
        )
        discipline_id = cursor.lastrowid
    cursor.execute(
        "INSERT OR IGNORE INTO program_disciplines (program_id, discipline_id, order_index) VALUES (?, ?, ?)",
        (program_id, discipline_id, _get_next_order_index(cursor, "program_disciplines", "program_id", program_id)),
    )
    return discipline_id


def _ensure_lesson_types(cursor) -> Dict[Optional[str], Optional[int]]:
    cursor.execute("SELECT id, name FROM lesson_types")
    existing = {row["name"]: row["id"] for row in cursor.fetchall()}
    for _, name in LESSON_TYPE_KEYWORDS:
        if name not in existing:
            cursor.execute("INSERT INTO lesson_types (name) VALUES (?)", (name,))
            existing[name] = cursor.lastrowid
    return existing


def _load_existing_topics(cursor, discipline_id: int) -> Dict[str, int]:
    cursor.execute(
        """
        SELECT t.id, t.title
        FROM topics t
        JOIN discipline_topics dt ON t.id = dt.topic_id
        WHERE dt.discipline_id = ?
        """,
        (discipline_id,),
    )
    return {_key(row["title"]): row["id"] for row in cursor.fetchall()}


def _load_existing_lessons(cursor, discipline_id: int) -> Dict[Tuple[int, str], Tuple[int, Dict[str, Optional[float]]]]:
    cursor.execute(
        """
        SELECT l.id, l.title, l.lesson_type_id, l.duration_hours,
               l.classroom_hours, l.self_study_hours, tl.topic_id
        FROM lessons l
        JOIN topic_lessons tl ON l.id = tl.lesson_id
        JOIN discipline_topics dt ON tl.topic_id = dt.topic_id
        WHERE dt.discipline_id = ?
        """,
        (discipline_id,),
    )
    results = {}
    for row in cursor.fetchall():
        results[(row["topic_id"], _key(row["title"]))] = (
            row["id"],
            {
                "lesson_type_id": row["lesson_type_id"],
                "duration_hours": row["duration_hours"],
                "classroom_hours": row["classroom_hours"],
                "self_study_hours": row["self_study_hours"],
            },
        )
    return results


def _update_lesson_if_missing(cursor, lesson_id: int, lesson: CurriculumLesson, lesson_types, existing):
    updates = {}
    if lesson.lesson_type_name and not existing.get("lesson_type_id"):
        updates["lesson_type_id"] = lesson_types.get(lesson.lesson_type_name)
    if lesson.total_hours is not None and not existing.get("duration_hours"):
        updates["duration_hours"] = lesson.total_hours
    if lesson.classroom_hours is not None and not existing.get("classroom_hours"):
        updates["classroom_hours"] = lesson.classroom_hours
    if lesson.self_study_hours is not None and not existing.get("self_study_hours"):
        updates["self_study_hours"] = lesson.self_study_hours
    if not updates:
        return
    set_clause = ", ".join(f"{key} = ?" for key in updates.keys())
    values = list(updates.values())
    values.append(lesson_id)
    cursor.execute(f"UPDATE lessons SET {set_clause} WHERE id = ?", values)


def _insert_questions(cursor, lesson_id: int, questions: List[CurriculumQuestion]) -> int:
    if not questions:
        return 0
    cursor.execute(
        """
        SELECT q.content
        FROM questions q
        JOIN lesson_questions lq ON q.id = lq.question_id
        WHERE lq.lesson_id = ?
        """,
        (lesson_id,),
    )
    existing = {_key(row["content"]) for row in cursor.fetchall()}
    added = 0
    for question in questions:
        content = _normalize_text(question.text)
        if _key(content) in existing:
            continue
        cursor.execute(
            """
            INSERT INTO questions (content, answer, difficulty_level, order_index)
            VALUES (?, ?, ?, ?)
            """,
            (content, "", 1, question.number or added + 1),
        )
        question_id = cursor.lastrowid
        cursor.execute(
            "INSERT OR IGNORE INTO lesson_questions (lesson_id, question_id, order_index) VALUES (?, ?, ?)",
            (lesson_id, question_id, question.number or added + 1),
        )
        existing.add(_key(content))
        added += 1
    return added


def _get_next_order_index(cursor, table: str, column: str, value: int) -> int:
    cursor.execute(f"SELECT MAX(order_index) as max_order FROM {table} WHERE {column} = ?", (value,))
    row = cursor.fetchone()
    max_order = row["max_order"] if row and row["max_order"] is not None else 0
    return max_order + 1
